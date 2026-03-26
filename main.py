import os
from openai import OpenAI
from dotenv import load_dotenv
import io
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
# Load environment variables from .env file
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)


def read_questions_from_file(file_path):
    """Extracts a list of questions (one per non-empty line) from a .txt, .docx, or .pdf file."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            raw_text = f.read()

    elif ext == ".docx":
        doc = Document(file_path)
        raw_text = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )

    elif ext == ".pdf":
        pages = extract_text_from_pdf(file_path)
        raw_text = "\n".join(page["text"] for page in pages)

    else:
        raise ValueError(f"Unsupported file type '{ext}'. Use .txt, .docx, or .pdf.")

    return [line.strip() for line in raw_text.splitlines() if line.strip()]


def new_answer_file_to_a_general_question_file(question_file):
    """Processes a batch of questions from a .txt, .docx, or .pdf file and saves answers to answers.txt."""
    try:
        question_list = read_questions_from_file(question_file)
        if not question_list:
            print("\n[Error] No questions found in the file.")
            return

        print(f"[Info] Found {len(question_list)} question(s). Processing...")
        answer_list = []
        for question in question_list:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": question}
                ],
                temperature=0.7
            )
            answer = response.choices[0].message.content
            answer_list.append(f"Q: {question}\nA: {answer}")

        with open('answers.txt', 'w') as answer_file:
            answer_file.write("\n\n".join(answer_list))
        print("\n[Success] Batch processing complete. Check 'answers.txt'.")
    except FileNotFoundError:
        print(f"\n[Error] The file '{question_file}' was not found.")
    except ValueError as e:
        print(f"\n[Error] {e}")


SCANNED_TEXT_THRESHOLD = 20  # Min characters per page to consider a PDF text-based


def is_scanned_pdf(pdf_path):
    """Returns True if the PDF is scanned (image-based), False if it contains real text."""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and len(text.strip()) > SCANNED_TEXT_THRESHOLD:
                return False
    return True


def extract_table_as_text(table):
    """Converts a pdfplumber table (list of rows) into a readable string."""
    if not table:
        return ""
    rows = []
    for row in table:
        row_cells = [str(cell).strip() if cell is not None else "" for cell in row]
        rows.append(" | ".join(row_cells))
    return "\n".join(rows)


def extract_image_text_from_page(page):
    """Runs OCR on images embedded inside a text-based PDF page (Hebrew + English)."""
    image_text_parts = []
    for img_obj in page.images:
        try:
            page_image = page.to_image(resolution=200).original
            bbox = (img_obj['x0'], img_obj['top'], img_obj['x1'], img_obj['bottom'])
            cropped = page_image.crop(bbox)
            ocr_text = pytesseract.image_to_string(cropped, lang='heb+eng').strip()
            if ocr_text:
                image_text_parts.append(ocr_text)
        except Exception:
            continue
    return "\n".join(image_text_parts)


def extract_text_from_text_pdf(pdf_path):
    """Extracts text, tables, and image-embedded text from a text-based PDF, page by page."""
    pages_data = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            parts = []

            # Regular text
            page_text = page.extract_text()
            if page_text and page_text.strip():
                parts.append(page_text.strip())

            # Tables
            tables = page.extract_tables()
            for table in tables:
                table_text = extract_table_as_text(table)
                if table_text:
                    parts.append(f"[TABLE]\n{table_text}\n[END TABLE]")

            # Text inside embedded images (OCR)
            if page.images:
                image_text = extract_image_text_from_page(page)
                if image_text:
                    parts.append(f"[IMAGE TEXT]\n{image_text}\n[END IMAGE TEXT]")

            combined_text = "\n\n".join(parts)
            if combined_text.strip():
                pages_data.append({
                    "text": combined_text.strip(),
                    "source_file": os.path.basename(pdf_path),
                    "page": page_num
                })
    return pages_data


def extract_text_from_scanned_pdf(pdf_path):
    """Uses OCR to extract text from all pages of a scanned PDF (supports Hebrew + English)."""
    pages_data = []
    images = convert_from_path(pdf_path, dpi=300)
    for page_num, image in enumerate(images, start=1):
        ocr_text = pytesseract.image_to_string(image, lang='heb+eng').strip()
        if ocr_text:
            pages_data.append({
                "text": ocr_text,
                "source_file": os.path.basename(pdf_path),
                "page": page_num
            })
    return pages_data


def extract_text_from_pdf(pdf_path):
    """Detects if a PDF is scanned or text-based and extracts its content accordingly.

    Returns a list of page dicts: [{"text": ..., "source_file": ..., "page": ...}, ...]
    Each dict represents one page and carries the metadata needed for future chunk referencing.
    """
    file_name = os.path.basename(pdf_path)
    if is_scanned_pdf(pdf_path):
        print(f"[INFO] '{file_name}' is a scanned PDF — using OCR.")
        return extract_text_from_scanned_pdf(pdf_path)
    else:
        print(f"[INFO] '{file_name}' is a text-based PDF — extracting directly.")
        return extract_text_from_text_pdf(pdf_path)


def process_pdf_files(pdf_paths):
    """Processes a list of PDF file paths and returns combined page-level extracted data.

    Each entry contains the text, source file name, and page number so that
    when the text is later split into chunks, every chunk can reference its origin.
    """
    all_pages = []
    for pdf_path in pdf_paths:
        print(f"\n[Processing] {os.path.basename(pdf_path)}")
        pages = extract_text_from_pdf(pdf_path)
        all_pages.extend(pages)
        print(f"[Done] Extracted {len(pages)} page(s) from '{os.path.basename(pdf_path)}'")
    print(f"\n[Summary] Total pages extracted: {len(all_pages)}")
    return all_pages


WORDS_PER_APPROX_PAGE = 300  # Word count used to simulate page breaks in DOCX files


def extract_docx_table_as_text(table_element):
    """Converts a DOCX table XML element into a readable string."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    rows = []
    for row_elem in table_element.findall(f"{{{W}}}tr"):
        cells = []
        for cell_elem in row_elem.findall(f"{{{W}}}tc"):
            cell_text = "".join(
                node.text for node in cell_elem.iter()
                if node.tag.endswith("}t") and node.text
            ).strip()
            cells.append(cell_text)
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_image_text_from_docx(doc):
    """Runs OCR on all images embedded in a DOCX file (Hebrew + English)."""
    image_texts = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                image = Image.open(io.BytesIO(image_data))
                ocr_text = pytesseract.image_to_string(image, lang='heb+eng').strip()
                if ocr_text:
                    image_texts.append(ocr_text)
            except Exception:
                continue
    return image_texts


def extract_text_from_docx(docx_path):
    """Extracts text, tables, and image-embedded text from a DOCX file.

    Since DOCX files have no native page numbers, content is grouped into approximate
    pages every WORDS_PER_APPROX_PAGE words. Returns the same list-of-dicts format as
    extract_text_from_pdf: [{"text": ..., "source_file": ..., "page": ...}, ...]
    """
    doc = Document(docx_path)
    pages_data = []
    current_parts = []
    current_word_count = 0
    page_num = 1

    def flush_page():
        nonlocal page_num, current_parts, current_word_count
        text = "\n\n".join(current_parts).strip()
        if text:
            pages_data.append({
                "text": text,
                "source_file": os.path.basename(docx_path),
                "page": page_num
            })
        page_num += 1
        current_parts = []
        current_word_count = 0

    # Iterate body elements in document order to preserve paragraph/table sequence
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            para_text = "".join(
                node.text for node in element.iter()
                if node.tag.endswith("}t") and node.text
            ).strip()
            if para_text:
                current_parts.append(para_text)
                current_word_count += len(para_text.split())
                if current_word_count >= WORDS_PER_APPROX_PAGE:
                    flush_page()

        elif tag == "tbl":
            table_text = extract_docx_table_as_text(element)
            if table_text.strip():
                current_parts.append(f"[TABLE]\n{table_text}\n[END TABLE]")
                current_word_count += len(table_text.split())
                if current_word_count >= WORDS_PER_APPROX_PAGE:
                    flush_page()

    # Images are stored in relationships, not inline with body elements
    image_texts = extract_image_text_from_docx(doc)
    if image_texts:
        current_parts.append("[IMAGE TEXT]\n" + "\n".join(image_texts) + "\n[END IMAGE TEXT]")

    # Flush any remaining content
    if current_parts:
        flush_page()

    return pages_data


def process_docx_files(docx_paths):
    """Processes a list of DOCX file paths and returns combined page-level extracted data.

    Each entry contains the text, source file name, and approximate page number so that
    when the text is later split into chunks, every chunk can reference its origin.
    """
    all_pages = []
    for docx_path in docx_paths:
        print(f"\n[Processing] {os.path.basename(docx_path)}")
        pages = extract_text_from_docx(docx_path)
        all_pages.extend(pages)
        print(f"[Done] Extracted {len(pages)} approx. page(s) from '{os.path.basename(docx_path)}'")
    print(f"\n[Summary] Total pages extracted: {len(all_pages)}")
    return all_pages

def get_chunks(pages_data, chunk_size, overlap):
    """Splits pages_data into overlapping chunks while preserving source metadata.

    Concatenates all page texts into one continuous string, tracks each page's
    start offset, then chunks with overlap. Each chunk is assigned the metadata
    (source_file, page) of the page where the chunk starts.

    Args:
        pages_data: list of dicts from process_pdf_files / process_docx_files,
                    each with keys "text", "source_file", "page".
        chunk_size: number of characters per chunk.
        overlap:    number of characters shared between consecutive chunks.

    Returns:
        list of dicts: [{"text": ..., "source_file": ..., "page": ...}, ...]
    """
    # Build one continuous string and record where each page starts
    full_text = ""
    page_offsets = []  # list of (start_char, source_file, page_number)
    for page in pages_data:
        page_offsets.append((len(full_text), page["source_file"], page["page"]))
        full_text += page["text"] + "\n\n"

    # Helper: find which page a character offset belongs to
    def get_metadata_at(char_index):
        metadata = page_offsets[0]
        for offset, source_file, page_num in page_offsets:
            if offset <= char_index:
                metadata = (offset, source_file, page_num)
            else:
                break
        return metadata[1], metadata[2]  # source_file, page_number

    chunk_list = []
    start = 0
    while start < len(full_text):
        chunk_text = full_text[start: start + chunk_size].strip()
        if chunk_text:
            source_file, page_num = get_metadata_at(start)
            chunk_list.append({
                "text": chunk_text,
                "source_file": source_file,
                "page": page_num
            })
        start += chunk_size - overlap
    return chunk_list


def create_vector_store(chunks_list):
    """Generates embeddings for all text chunks and preserves their source metadata."""
    response = client.embeddings.create(
        input=[chunk["text"] for chunk in chunks_list],
        model="text-embedding-3-small"
    )

    vector_store = []
    for i, item in enumerate(response.data):
        vector_store.append({
            "text": chunks_list[i]["text"],
            "source_file": chunks_list[i]["source_file"],
            "page": chunks_list[i]["page"],
            "vector": item.embedding
        })
    return vector_store


def retrieve_k_relevant_chunks(user_query, database, k_value):
    """Finds top-K most relevant chunks based on vector similarity."""
    response = client.embeddings.create(
        input=user_query,
        model="text-embedding-3-small"
    )
    query_vector = response.data[0].embedding

    results = []
    for item in database:
        # Calculate dot product similarity
        score = sum(a * b for a, b in zip(query_vector, item["vector"]))
        results.append({
            "text": item["text"],
            "source_file": item["source_file"],
            "page": item["page"],
            "score": score
        })

    # Sort results by similarity score
    results.sort(key=lambda x: x["score"], reverse=True)
    return [{"text": r["text"], "source_file": r["source_file"], "page": r["page"]} for r in results[:k_value]]


def generate_answer(user_query, context):
    """Sends the retrieved context to GPT-4o for a response."""
    prompt = f"""
    Answer the following question using ONLY the provided context.
    If the answer is not in the context, your response MUST be exactly: "i don't have enough data"

    Context:
    {context}

    Question:
    {user_query}
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that strictly follows the provided context."},
            {"role": "user", "content": prompt}
        ],
        temperature=0
    )
    return response.choices[0].message.content


def collect_files_from_user():
    """Prompts the user to enter file paths one by one until they are done.

    Validates that each path exists and has a supported extension (.pdf or .docx).
    Returns two lists: pdf_paths and docx_paths.
    """
    pdf_paths = []
    docx_paths = []
    print("\n--- File Upload ---")
    print("Enter the full path to each file you want to upload.")
    print("Supported formats: .pdf, .docx  |  Type 'done' when finished.\n")

    while True:
        raw = input("File path: ").strip().strip('"').strip("'")
        if raw.lower() == "done":
            if not pdf_paths and not docx_paths:
                print("[Error] You must upload at least one file.")
                continue
            break
        if not raw:
            continue
        if not os.path.isfile(raw):
            print(f"[Error] File not found: {raw}")
            continue
        ext = os.path.splitext(raw)[1].lower()
        if ext == ".pdf":
            pdf_paths.append(raw)
            print(f"[Added] {os.path.basename(raw)}")
        elif ext == ".docx":
            docx_paths.append(raw)
            print(f"[Added] {os.path.basename(raw)}")
        else:
            print(f"[Error] Unsupported file type '{ext}'. Only .pdf and .docx are supported.")

    return pdf_paths, docx_paths


def main():
    print("--- Welcome to the Advanced RAG & Q&A Tool ---")
    print("\nAvailable Modes:")
    print("1. Interactive RAG Chat")
    print("2. Batch Q&A Processor")

    user_choice = input("\nChoose mode (1 or 2): ").strip()
    if user_choice == "1":
        # --- RAG Mode Execution ---
        try:
            print("\n--- RAG Setup ---")
            c_size = input("Enter Chunk Size [default 500]: ") or "500"
            c_overlap = input("Enter Overlap [default 50]: ") or "50"

            chunk_size = int(c_size.strip())
            overlap = int(c_overlap.strip())
            k_value = 1  # Starting with K=1 to demonstrate the system's ability to adjust

            # Collect files from user and extract text with source metadata
            pdf_paths, docx_paths = collect_files_from_user()
            pages_data = process_pdf_files(pdf_paths) + process_docx_files(docx_paths)

            print(f"\nInitializing RAG (Size: {chunk_size}, Overlap: {overlap})...")
            chunks = get_chunks(pages_data, chunk_size, overlap)
            database = create_vector_store(chunks)
            print("\n[System] RAG Mode Activated. Type 'exit' to quit.")

            while True:
                user_query = input("\nUser: ")
                if user_query.lower() == 'exit':
                    break
                if not user_query.strip():
                    continue

                # Retrieval Step
                relevant_chunks = retrieve_k_relevant_chunks(user_query, database, k_value)
                context = "\n\n---\n\n".join(chunk["text"] for chunk in relevant_chunks)

                # Generation Step
                answer = generate_answer(user_query, context)
                print(f"\nAI (K={k_value}): {answer}")

                # Show source references
                references = [f"{c['source_file']} (page {c['page']})" for c in relevant_chunks]
                print(f"[Sources] {', '.join(references)}")

                # Smart Logic: Offer to increase K if context is insufficient
                if answer.lower() == "i don't have enough data":
                    print("\n[System Alert] The information was not found in the current context.")
                    suggestion = input("Would you like to try again with a larger K-value? (yes/no): ")

                    if suggestion.lower().strip() == "yes":
                        try:
                            new_k = input("Enter new K-value (e.g., 3 or 5): ")
                            k_value = int(new_k)
                            print(f"[System] K-value updated to {k_value}. Please ask your question again.")
                        except ValueError:
                            print("[Error] Invalid number. Keeping current K.")
                    else:
                        exit_confirm = input("Would you like to exit the program? (yes/no): ")
                        if exit_confirm.lower().strip() == "yes":
                            break

                print("-" * 50)

        except Exception as e:
            print(f"\n[Error] An error occurred during setup: {e}")

    elif user_choice == "2":
        # --- Batch Mode Execution ---
        print("\n--- Batch Mode Activated ---")
        print("Supported formats: .txt, .docx, .pdf")
        question_file = input("Enter path to your questions file: ").strip().strip('"').strip("'")
        if not os.path.isfile(question_file):
            print(f"\n[Error] File not found: {question_file}")
        else:
            new_answer_file_to_a_general_question_file(question_file)

    else:
        print("\n[System] Invalid choice. Program terminated.")


if __name__ == "__main__":
    main()

